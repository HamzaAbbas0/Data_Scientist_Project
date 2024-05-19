from IPython.display import display
from ipywidgets import widgets
import json
from databrickschatbotapi.DatascientistPipeline.report_gen import AnalysisGenerator
from databrickschatbotapi.DatascientistPipeline.Graph_description import GraphDescriptionPipeline
import json
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import time
import os
import torch

# if torch.cuda.is_available():
#     torch.cuda.init()
class DataScienceReport:
    def __init__(self, auth_token, df, question,json_file, problem_type,process_id, source, file_name):
        if torch.cuda.is_available():
            torch.cuda.init()
        self.problem_type = problem_type
        self.source = source
        self.file_name = file_name
        self.question = question
        self.json_file = json_file
        self.process_id = process_id
#         self.cross_tab1 = cross_tab1
#         self.cross_tab2 = cross_tab2
#         self.rep_instance = AnalysisGenerator(auth_token, df, question)
#         self.graph_function = GraphDescriptionPipeline()
#         self.dependent_variable = dependent_variable
        self.json_output_file = 'latest_output.json'
        self.doc_output_file = 'Dynamic Data Science Report.docx'
        self.report_sections = [{
        "names": "From Data to Insight: A Comprehensive Data Science Exploration Report"
        }]

#         self.intro_check = widgets.Checkbox(False, description='Write Introduction')
# #         self.corrmat_check = widgets.Checkbox(False, description='Write Correlation Analysis?')
#         self.domain_qn_ans_check = widgets.Checkbox(False, description='Write Domain Question Answer?')
#         self.summary_result_check = widgets.Checkbox(False, description='Write Summary Statistics?')
#         self.most_corrmat_check = widgets.Checkbox(False, description='Write Most Correlation Features Result?')
#         self.most_imp_featur_check = widgets.Checkbox(False, description='Write Most Important Features?')
        
#         self.cross_tab_check =  widgets.Checkbox(False, description='Find Cross Tab?')
#         self.chi_square_check =  widgets.Checkbox(False, description='Find Find Chi Square?')
        
#         self.graph_to_text_check=  widgets.Checkbox(False, description='Graph To Text Analysis?')
#         self.graph_distribution = widgets.Checkbox(False, description='Distributions')
#         self.graph_lstm_inference = widgets.Checkbox(False, description='LSTM_inference_plot')
#         self.graph_training_loss = widgets.Checkbox(False, description='Training_loss_graph')
#         self.graph_heatmap_explainer = widgets.Checkbox(False, description='Heatmap_explainer')
#         self.graph_missing_number= widgets.Checkbox(False, description='Missing_Number')
#         self.graph_most_correlated_features = widgets.Checkbox(False, description='Most_Correlated_Features')
#         self.graph_trend_graph = widgets.Checkbox(False, description='Trend_Graph')
#         self.graph_barchart = widgets.Checkbox(False, description='BarChart')
#         self.graph_histogram = widgets.Checkbox(False, description='Histogram')
#         self.graph_cross_tabulation = widgets.Checkbox(False, description='CrossTabulation')
#         self.graph_pairwise = widgets.Checkbox(False, description='PairWise')

        self.intro_check = False
#         self.corrmat_check = widgets.Checkbox(False, description='Write Correlation Analysis?')
        self.domain_qn_ans_check = False
        self.summary_result_check = False
        self.most_corrmat_check = False
        self.most_imp_featur_check = False
        self.confusion_matrix_check = False
        self.cross_tab_check =  False
        self.chi_square_check =  False
        
        self.graph_to_text_check=  False
        self.graph_distribution = False
        self.graph_lstm_inference = False
        self.graph_training_loss = False
        self.graph_heatmap_explainer = False
        self.graph_missing_number= False
        self.graph_most_correlated_features = False
        self.graph_trend_graph = False
        self.graph_barchart = False
        self.graph_histogram = False
        self.graph_cross_tabulation = False
        self.graph_pairwise = False
        self.confusion_matrix_graph = False
        self.Neural_Regressor = False
        
        print(f"the problem Type is : {problem_type} ")
        if problem_type.lower() == "time series":
#         if next(iter(json_file.values())).split('/')[0] == 'Time Series':
            self.display_widgets_timeseries()
        elif problem_type.lower() == "numerical":
#         elif next(iter(json_file.values())).split('/')[0] == 'Numerical':
            self.display_widgets_numerical()
        elif problem_type.lower() == "categorical":
#         elif next(iter(json_file.values())).split('/')[0] == 'Categorical':
            self.display_widgets_catagorical()
        else:
            print("Something went wrong plz check the problem type existing path")
            

    def display_widgets_timeseries(self):
#         print('i am the time series checkbox')
#         display(self.intro_check)
#         display(self.summary_result_check)
# #         display(self.corrmat_check)
# #         display(self.domain_qn_ans_check)
#         display(self.most_corrmat_check)
#         display(self.most_imp_featur_check)
        
# #         display(self.graph_to_text_check)
#         display(self.graph_distribution)
#         display(self.graph_lstm_inference)
#         display(self.graph_training_loss)
#         display(self.graph_heatmap_explainer)
#         display(self.graph_missing_number)
#         display(self.graph_most_correlated_features)
#         display(self.graph_trend_graph)
        
#         print('i am the time series checkbox')
        self.intro_check = True
        self.summary_result_check = True
#         display(self.corrmat_check)
#         display(self.domain_qn_ans_check)
        self.most_corrmat_check = True
        self.most_imp_featur_check= True
        
#         display(self.graph_to_text_check)
        self.graph_distribution = True
        self.graph_lstm_inference = True
        self.graph_training_loss = True
        self.graph_heatmap_explainer = True
        self.graph_missing_number= True
        self.graph_most_correlated_features= True
        self.graph_trend_graph = True
        
    def display_widgets_numerical(self):
#         print('i am the numerical columns checkbox')
#         display(self.intro_check)
# #         display(self.corrmat_check)
# #         display(self.domain_qn_ans_check)
#         display(self.summary_result_check)
#         display(self.most_corrmat_check)
# #         display(self.most_imp_featur_check)
        
# #         display(self.graph_to_text_check)
#         display(self.graph_barchart)
#         display(self.graph_histogram)
#         display(self.graph_cross_tabulation)
#         display(self.graph_pairwise)

#         print('i am the numerical columns checkbox')
        self.intro_check = True
#         display(self.corrmat_check)
#         display(self.domain_qn_ans_check)
        self.summary_result_check = True
        self.most_corrmat_check = True
#         display(self.most_imp_featur_check)
        
#         display(self.graph_to_text_check)
#         self.graph_barchart = True
#         self.graph_histogram = True
#         self.graph_cross_tabulation = True
        self.graph_pairwise = True
        self.graph_heatmap_explainer = True
        self.graph_missing_number= True
        self.graph_distribution = True
#         self.graph_lstm_inference = True
        self.Neural_Regressor = True
        
    def display_widgets_catagorical(self):
#         print('i am the Catagorical columns checkbox')
#         #display(self.cross_tab_check)
#         display(self.intro_check)
# #         display(self.corrmat_check)
# #         display(self.domain_qn_ans_check)
# #         display(self.summary_result_check)
# #         display(self.most_corrmat_check)
# #         display(self.most_imp_featur_check)
#         display(self.most_corrmat_check)
#         display(self.cross_tab_check)
#         display(self.chi_square_check)
        
# #         display(self.graph_to_text_check)
#         display(self.graph_barchart)
#         display(self.graph_histogram)
#         display(self.graph_cross_tabulation)
        
#         print('i am the Catagorical columns checkbox')
        #display(self.cross_tab_check)
        self.intro_check = True
#         display(self.corrmat_check)
#         display(self.domain_qn_ans_check)
#         display(self.summary_result_check)
#         display(self.most_corrmat_check)
#         display(self.most_imp_featur_check)
        self.most_corrmat_check = True
        self.cross_tab_check = True
        self.chi_square_check = True
        self.confusion_matrix_check = True
#         display(self.graph_to_text_check)
#         self.graph_barchart = True
#         self.graph_histogram = True
#         self.graph_cross_tabulation = True
        self.graph_heatmap_explainer = True
        self.graph_missing_number= True
        self.confusion_matrix_graph = True
        self.graph_distribution = True
        

    def add_section(self, section_name, variable_name):
        section = {"names": section_name, "variable_name": variable_name}
        self.report_sections.append(section)
        
    def add_section_image(self, section_name, image_path, graph_result):
        section = {"names":section_name,"image_path":image_path,"image_variable":graph_result}
        self.report_sections.append(section)
        
    def add_image(self, doc, path, width=Inches(5), height=Inches(5)):
        doc.add_picture(path, width, height)
        
    def text_analysis(self, auth_token, df, question):
        self.rep_instance = AnalysisGenerator(auth_token, df, question)
        try:
            if self.intro_check:
                intro = self.rep_instance.introduction()
                self.add_section("Introduction", intro)
        except Exception as e:
            print(f"An error occurred in Introduction section: {e}")

#         if self.corrmat_check.value :
#             corrmat_path = f'{self.problem_type}/csv/{self.file_name}_correlation_matrix.csv'
#             corrmat_result = self.rep_instance.analyze_csv_and_corrmat(corrmat_path,self.dependent_variable)
#             self.add_section("Correlation Analysis", corrmat_result,)

        try:
            if self.domain_qn_ans_check:
                question = 'brief me the education details of education columns'
                domain_qn_ans = self.rep_instance.analyze_user_question(question, user_specified_domain="any")
                self.add_section("Domain Question Answer", domain_qn_ans)
        except Exception as e:
            print(f"An error occurred in Domain Question Answer section: {e}")

        try:
            if self.summary_result_check:
                key_to_extract = 'Summary Statistics '
                value = self.json_file[key_to_extract]
    #             question = 'write me the overall csv file description in summary wise'
    #             csv_file_path = f'{self.problem_type}/csv/{self.file_name}_summary_statistics.csv'
                summary_result = self.rep_instance.description(value)
                self.add_section("Summary Statistics", summary_result)
        except Exception as e:
            print(f"An error occurred in Summary Statistics section: {e}")
        
        try:
            if self.confusion_matrix_check:
                key_to_extract = 'Random Forest Classifier confusion matrix csv path '
                value = self.json_file[key_to_extract]

                confusion_result = self.rep_instance.confusion_matrix(value)
                self.add_section("Confusion-Matrix", confusion_result)
        except Exception as e:
            print(f"An error occurred in Confusion Matrix section: {e}")
            
        try:
            if self.confusion_matrix_check:
                key_to_extract = 'All correlated Features with dependent Variable '
                value = self.json_file[key_to_extract]

                confusion_result = self.rep_instance.confusion_matrix(value)
                self.add_section("Confusion-Matrix", confusion_result)
        except Exception as e:
            print(f"An error occurred in Confusion Matrix section: {e}")

        try:
            if self.most_corrmat_check:
                key_to_extract = 'Most correlated Features with dependent Variable '
                value = self.json_file[key_to_extract]
                most_corrmat_result = self.rep_instance.analyze_csv_and_mostcorrmat(value)
                self.add_section("Most Co-Relation Features", most_corrmat_result)
        except Exception as e:
            print(f"An error occurred in Most Co-Relation Features section: {e}")


#         if self.most_imp_featur_check:
#             csv_file_path = f'{self.problem_type}/csv/{self.file_name}_most_correlated_features.csv'
#             most_imp_featur = self.rep_instance.analyze_imp_features(csv_file_path)
#             self.add_section("Most Important Features", most_imp_featur)
            
# #             try:
# #                 if self.graph_most_correlated_features.value:
#                     #query = 'explain heatmap graph with deep insight facts'
#                     img_file_path = f'{self.problem_type}/graphs/{self.file_name}_most_correlated_features.png'
#                     graph_result = self.graph_function.Most_Correlated_Features(img_file_path, query)
#                     self.add_section_image("Graph To Text Generation",img_file_path,graph_result)
#             except Exception as e:
#                 print(f"An error occurred: {e}")
            
        try:
            if self.chi_square_check:
                key_to_extract = 'Chi-Square statistics '
                value = self.json_file[key_to_extract]
                chi_square_result = self.rep_instance.chi_square(value) 
                self.add_section("Chi Square Statistics", chi_square_result)
        except Exception as e:
            print(f"An error occurred in Chi Square Statistics section: {e}")

#             print(chi_square_result)
            
#         if self.cross_tab_check:
# #             cross_tab1 = 'city'
# #             cross_tab2 = 'deal_size'
#             cross_tab1 = input("write column1 for the cross tab")
#             cross_tab2 = input("write column2 for the cross tab")
        
        
# #             if os.path.isfile(f'{self.problem_type}/csv/{self.file_name}{cross_tab1}_vs_{cross_tab2}.csv'):
# #                 full_path = f'{self.problem_type}/csv/{self.file_name}{cross_tab1}_vs_{cross_tab2}.csv'
# #             elif os.path.isfile(f'{self.problem_type}/csv/{self.file_name}{cross_tab2}_vs_{cross_tab1}.csv'):
# #                 full_path = f'{self.problem_type}/csv/{self.file_name}{cross_tab2}_vs_{cross_tab1}.csv'
# #             else:
# #                 full_path = None
#             key_to_extract = 'Crosstabs of two columns '
#             value = self.json_file[key_to_extract]    
               
#             csv_file_path = f'{value}_{cross_tab1}_vs_{cross_tab2}.csv' 
# #             question = 'write me the cross tab of the city and small'    
#             cross_tab_result = self.rep_instance.cross_tab(csv_file_path,cross_tab1,cross_tab2)
#             self.add_section("Cross Tab", cross_tab_result)







        ###### Graph Expanation ########
    def graph_analyze(self, question):
        self.graph_function = GraphDescriptionPipeline()

        try:
            if self.graph_to_text_check:
                image_path = "Time Series-copy/graphs/denormalized_denormalized_dataset_new_LSTM_inference_plot_['Mold'].png"
                graph_result = self.graph_function.LSTM_inference_plot(image_path, self.question)
                self.add_section_image("Graph To Text Generation", image_path, graph_result)
        except Exception as e:
            print(f"An error occurred in Graph To Text Generation section: {e}")

        try:
            if self.graph_most_correlated_features:
                json_value = self.json_file['Graph of most correlated Features with dependent Variable ']
                graph_result = self.graph_function.Most_Correlated_Features(json_value, self.question)
                self.add_section_image("Most Correlated Feature Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Most Correlated Feature Graph Analysis section: {e}")

        try:
            if self.graph_distribution:
                json_value = self.json_file['Probability distributions ']
                graph_result = self.graph_function.Distributions(json_value, self.question)
                self.add_section_image("Distribution Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Distribution Graph Analysis section: {e}")

        try:
            if self.graph_lstm_inference:
                json_value = self.json_file['LSTM Inference plot ']
                graph_result = self.graph_function.LSTM_inference_plot(json_value, self.question)
                self.add_section_image("LSTM Inference Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in LSTM Inference Graph Analysis section: {e}")

        try:
            if self.graph_training_loss:
                json_value = self.json_file['LSTM Training loss graph ']
                graph_result = self.graph_function.Training_loss_graph(json_value, self.question)
                self.add_section_image("Training Loss Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Training Loss Graph Analysis section: {e}")

        try:
            if self.graph_pairwise:
                json_value = self.json_file['Pairwise plots ']
                graph_result = self.graph_function.PairWise_Graph(json_value, self.question)
                self.add_section_image("PairWise Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in PairWise Graph Analysis section: {e}")

        try:
            if self.graph_missing_number:
                json_value = self.json_file['Missing number plot before cleaning ']
                graph_result = self.graph_function.Missing_Number(json_value, self.question)
                self.add_section_image("Missing Numbers Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Missing Numbers Graph Analysis section: {e}")

        try:
            if self.graph_heatmap_explainer:
                json_value = self.json_file['Correlation Heatmap ']
                graph_result = self.graph_function.Heatmap_explainer(json_value, self.question)
                self.add_section_image("Heat_Explainer Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Heat_Explainer Graph Analysis section: {e}")

        try:
            if self.confusion_matrix_graph:
                json_value = self.json_file['Random Forest Classifier confusion matrix graph path ']
                graph_result = self.graph_function.Confusion_matrix(json_value, self.question)
                self.add_section_image("Confusion_matrix Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Confusion_matrix Graph Analysis section: {e}")

        try:
            if self.graph_trend_graph:
                json_value = self.json_file['Trend graph of denormalized_GlobalLandTemperaturesByMajorCity ']
                graph_result = self.graph_function.Trend_Graph(json_value, self.question)
                self.add_section_image("Trend Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Trend Graph Analysis section: {e}")

        try:
            if self.Neural_Regressor:
                json_value = self.json_file['Multi-linear Regression Inference plot ']
                ImageShow = self.graph_function.display_image(json_value)
                print(ImageShow)
                graph_result = self.graph_function.Neural_Regressor(json_value, self.question)
                self.add_section_image("Multi-linear Regression Inference Graph Analysis", json_value, graph_result)
        except Exception as e:
            print(f"An error occurred in Multi-linear Regression Inference Graph Analysis section: {e}")




#         if self.graph_cross_tabulation:
#              json_value=self.json_file['Cross tabulation graph of two columns ']
#             var1 = input("write  first column  for Cross tabulation Graph visualize  : ")
#             var2 = input("write Second column for Cross tabulation Graph visualize   : ")

#             img_file_path = f'{self.problem_type}/graphs/{self.file_name}_Cross_tabulation_{var1}_vs_{var2}.png'
#             graph_result = self.graph_function.CrossTabulation(img_file_path,self.question,var1,var2)
#             self.add_section_image("Cross Tabulation Graph Analysis",img_file_path,graph_result)



#         if self.graph_barchart:
#              json_value=self.json_file['Bar charts ']

#             var1 = input("write column for bar chart visualize     : ")

#             img_file_path = f'{self.problem_type}/graphs/{self.file_name}_bar_chart_{var1}.png'
#             graph_result = self.graph_function.BarCharts(img_file_path, self.question,var1)
#             self.add_section_image("Bar Chart Graph Analysis",img_file_path,graph_result)


#         if self.graph_histogram.value:
#              json_value=self.json_file['Histogram of two columns ']
# #             var1 = input("write  first column  for Histogram visualize  : ")
# #             var2 = input("write Second column for Histogram visualize   : ")

# #             img_file_path = f'{self.problem_type}/graphs/{self.file_name}_histogram_{var1}_vs_{var2}.png'
#             graph_result = self.graph_function.Histogram(img_file_path,self.question,var1,var2)
#             self.add_section_image("Histogram Graph Analysis ",img_file_path,graph_result)

        # Generate Word document
#         self.generate_word_document()
        
    def generate_word_document(self):
        try:
            # Save the JSON to a file
            output_file_path = f'Knowledge/{self.problem_type}/{self.source}/{self.process_id}/report/{self.file_name}.json'
            with open(output_file_path, 'w') as json_file:
                json.dump({"sections": self.report_sections}, json_file, indent=2)

            print(f"JSON data has been saved to {output_file_path}")

            doc = Document()
            json_path = f'Knowledge/{self.problem_type}/{self.source}/{self.process_id}/report/{self.file_name}.json'
            with open(json_path, 'r') as json_file:
                data = json.load(json_file)

            for section in data["sections"]:
                doc.add_heading(section["names"], level=0)

                if "variable_name" in section:
                    doc.add_paragraph(section["variable_name"])

                if "image_path" and "image_variable" in section:
                    self.add_image(doc, section["image_path"])
                    doc.add_paragraph(section["image_variable"])
            doc_path = f'Knowledge/{self.problem_type}/{self.source}/{self.process_id}/report/{self.file_name}.docx'
            doc.save(doc_path)
            
            print(f'Report is saved at: {doc_path}')
            return doc_path
            

        except Exception as e:
            print(f"An error occurred while generating the Word document: {e}")

        
        



